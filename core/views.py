from django.shortcuts import render, get_object_or_404
from django.contrib.auth.models import User
from rest_framework import generics, permissions, status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from .serializers import DeCuongSerializer, HocVienSerializer
from .models import HocVien, DeCuong
import google.generativeai as genai
import os
from io import BytesIO
from django.http import HttpResponse

# Import cho Word
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. Đăng ký User ---
@api_view(['POST'])
@permission_classes([permissions.AllowAny])
def register_user(request):
    username = request.data.get('username')
    password = request.data.get('password')
    ho_ten = request.data.get('ho_ten')

    if not username or not password or not ho_ten:
        return Response({'error': 'Vui lòng cung cấp đầy đủ thông tin'}, status=status.HTTP_400_BAD_REQUEST)

    if User.objects.filter(username=username).exists():
        return Response({'username': 'Tên đăng nhập đã tồn tại.'}, status=status.HTTP_400_BAD_REQUEST)

    # Tạo User
    user = User.objects.create_user(username=username, password=password)
    # Tạo HocVien liên kết
    HocVien.objects.create(user=user, ho_ten=ho_ten)
    
    return Response({'success': 'Tạo tài khoản thành công'}, status=status.HTTP_201_CREATED)

# --- 2. Xem và Cập nhật Đề cương ---
@api_view(['GET', 'POST'])
@permission_classes([permissions.IsAuthenticated])
def de_cuong_view(request):
    hoc_vien = get_object_or_404(HocVien, user=request.user)
    
    # Tạo hoặc lấy đề cương
    de_cuong, created = DeCuong.objects.get_or_create(hoc_vien=hoc_vien)

    if request.method == 'POST':
        # Cập nhật đề cương
        
        # Chỉ cho phép sửa nếu trạng thái không phải là 'đã nộp' hoặc 'đã duyệt'
        if de_cuong.trang_thai in ['da_nop', 'da_duyet']:
             return Response({'error': 'Đề cương đã nộp, không thể chỉnh sửa.'}, status=status.HTTP_403_FORBIDDEN)
        
        serializer = DeCuongSerializer(de_cuong, data=request.data, partial=True) # partial=True cho phép cập nhật 1 phần
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    else:
        # (GET) Lấy đề cương
        serializer = DeCuongSerializer(de_cuong)
        return Response(serializer.data)

# --- 3. Gợi ý AI ---
@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def goi_y_ai(request):
    try:
        genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
        model = genai.GenerativeModel('models/gemini-pro-latest')
    except Exception as e:
        print(f"================ LỖI CẤU HÌNH GOOGLE AI ================")
        print(e)
        return Response({'error': 'Lỗi cấu hình AI trên máy chủ.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    prompt = request.data.get('prompt')
    if not prompt:
        return Response({'error': 'Không có prompt'}, status=status.HTTP_400_BAD_REQUEST)

    # Tạo prompt hoàn chỉnh
    full_prompt = f"""
    Bạn là một trợ lý học thuật, đang giúp một học viên cao học viết đề cương luận văn.
    Học viên đã viết phần sau:
    ---
    {prompt}
    ---
    Dựa trên nội dung trên, hãy đưa ra gợi ý chuyên sâu (khoảng 100-150 từ) để giúp họ:
    1. Phát triển ý tưởng rõ ràng hơn.
    2. Chỉ ra các luận điểm còn thiếu.
    3. Đề xuất các bước tiếp theo.
    Hãy trả lời bằng tiếng Việt.
    """
    
    try:
        response = model.generate_content(full_prompt)
        # Bỏ dấu * và định dạng lại
        clean_text = response.text.replace('*', '').strip()
        return Response({'goi_y': clean_text})
    except Exception as e:
        print(f"================ LỖI GOOGLE AI ================")
        print(e)
        return Response({'error': f'Lỗi từ Google AI: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# --- 4. Tải Word (.docx) ---
@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def download_pdf_view(request):
    try:
        # 1. Lấy dữ liệu
        hoc_vien = get_object_or_404(HocVien, user=request.user)
        de_cuong = get_object_or_404(DeCuong, hoc_vien=hoc_vien)

        # 2. Tạo tài liệu Word
        document = Document()
        
        # 3. Thêm nội dung
        
        # Tiêu đề (Sửa lỗi cú pháp: đảm bảo các dòng riêng biệt)
        title = document.add_heading(de_cuong.ten_de_tai or "(Chưa có tên đề tài)", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        
        # Thông tin học viên
        document.add_heading('Thông tin học viên', level=2)
        p_info = document.add_paragraph()
        p_info.add_run('Họ và tên: ').bold = True
        p_info.add_run(hoc_vien.ho_ten)
        
        p_info2 = document.add_paragraph()
        p_info2.add_run('Ngày sinh: ').bold = True
        p_info2.add_run(str(hoc_vien.ngay_sinh) if hoc_vien.ngay_sinh else "(Chưa cập nhật)")
        
        p_info3 = document.add_paragraph()
        p_info3.add_run('Quê quán: ').bold = True
        p_info3.add_run(hoc_vien.que_quan or "(Chưa cập nhật)")

        # Các chương (Sửa lỗi đánh số thứ tự)
        sections = [
            ('1. Lý do chọn đề tài (Tính cấp thiết)', de_cuong.ly_do_chon_de_tai),
            ('2. Chương 1: Khung lý thuyết', de_cuong.khung_ly_thuyet),
            ('3. Chương 2: Thiết kế và Tổ chức', de_cuong.thiet_ke_va_to_chuc),
            ('4. Chương 3: Thực nghiệm', de_cuong.thuc_nghiem),
        ]

        for title, content in sections:
            document.add_heading(title, level=2)
            document.add_paragraph(content or "(Chưa có nội dung)")

        # 4. Lưu vào buffer
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # 5. Tạo HttpResponse
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="DeCuongLuanVan.docx"'
        return response

    except Exception as e:
        print(f"Lỗi tạo Word: {e}")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# --- 5. View Trang chủ (index.html) ---
# View này phục vụ tệp index.html
def index(request):
    return render(request, 'core/index.html')